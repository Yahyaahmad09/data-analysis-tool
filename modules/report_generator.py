"""
Report Generator Module
Generate PDF, HTML, and DOCX reports with customization
"""

import io
import json
from datetime import datetime
from typing import Dict, List, Optional
import pandas as pd

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ReportGenerator:
    """Generate comprehensive analysis reports"""

    def __init__(self, template: str = 'professional', customizations: Optional[Dict] = None):
        self.template = template
        self.customizations = customizations or {}
        self.styles = self._get_styles()

    def _get_styles(self) -> Dict:
        """Get report styles based on template"""
        base_styles = {
            'font': 'Helvetica',
            'primary_color': '#0072B2',
            'secondary_color': '#009E73',
            'accent_color': '#D55E00'
        }

        if self.template == 'professional':
            base_styles.update({
                'title_size': 24,
                'header_size': 18,
                'body_size': 11
            })
        elif self.template == 'academic':
            base_styles.update({
                'title_size': 20,
                'header_size': 16,
                'body_size': 10
            })
        elif self.template == 'minimal':
            base_styles.update({
                'title_size': 22,
                'header_size': 16,
                'body_size': 11
            })
        elif self.template == 'detailed':
            base_styles.update({
                'title_size': 26,
                'header_size': 20,
                'body_size': 12
            })

        return base_styles

    def generate_html_report(self, analysis_data: Dict) -> str:
        """Generate HTML report"""
        report_title = self.customizations.get('title', 'Data Analysis Report')
        report_notes = self.customizations.get('notes', '')

        # Build HTML
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report_title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 2rem;
            background: #fafafa;
        }}
        .header {{
            background: linear-gradient(90deg, {self.styles['primary_color']} 0%, {self.styles['secondary_color']} 100%);
            color: white;
            padding: 2rem;
            border-radius: 12px;
            margin-bottom: 2rem;
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }}
        .header h1 {{
            margin: 0;
            font-size: {self.styles['title_size']}px;
        }}
        .header .meta {{
            margin-top: 1rem;
            opacity: 0.9;
            font-size: 0.9rem;
        }}
        .section {{
            background: white;
            padding: 2rem;
            margin-bottom: 1.5rem;
            border-radius: 12px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        }}
        .section h2 {{
            color: {self.styles['primary_color']};
            font-size: {self.styles['header_size']}px;
            border-bottom: 2px solid {self.styles['primary_color']};
            padding-bottom: 0.5rem;
            margin-top: 0;
        }}
        .section h3 {{
            color: {self.styles['secondary_color']};
            font-size: 1.3rem;
            margin-top: 1.5rem;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1rem 0;
            font-size: 0.9rem;
        }}
        th {{
            background: {self.styles['primary_color']};
            color: white;
            padding: 0.75rem;
            text-align: left;
            font-weight: 600;
        }}
        td {{
            padding: 0.75rem;
            border-bottom: 1px solid #e0e0e0;
        }}
        tr:hover {{
            background: #f5f5f5;
        }}
        .metric-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1rem;
            margin: 1rem 0;
        }}
        .metric-card {{
            background: linear-gradient(135deg, #ffffff 0%, #f0f8ff 100%);
            padding: 1.25rem;
            border-radius: 8px;
            border-left: 4px solid {self.styles['primary_color']};
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }}
        .metric-label {{
            font-size: 0.85rem;
            color: #666;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .metric-value {{
            font-size: 1.8rem;
            font-weight: 700;
            color: {self.styles['primary_color']};
            margin-top: 0.25rem;
        }}
        .quality-score {{
            display: inline-block;
            background: {self.styles['secondary_color']};
            color: white;
            padding: 0.5rem 1.5rem;
            border-radius: 24px;
            font-size: 1.2rem;
            font-weight: 700;
        }}
        .footer {{
            text-align: center;
            color: #888;
            font-size: 0.85rem;
            margin-top: 3rem;
            padding-top: 1.5rem;
            border-top: 1px solid #e0e0e0;
        }}
        .alert {{
            padding: 1rem;
            border-radius: 8px;
            margin: 1rem 0;
        }}
        .alert-info {{ background: #e6f7ff; border-left: 4px solid #1890ff; }}
        .alert-warning {{ background: #fffbe6; border-left: 4px solid #faad14; }}
        .alert-success {{ background: #f6ffed; border-left: 4px solid #52c41a; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{report_title}</h1>
        <div class="meta">
            <strong>Generated:</strong> {datetime.now().strftime('%B %d, %Y at %H:%M')}<br>
            <strong>File:</strong> {analysis_data.get('file_name', 'N/A')}<br>
            {f'<strong>Notes:</strong> {report_notes}' if report_notes else ''}
        </div>
    </div>
"""

        # Table of Contents
        sections = self.customizations.get('sections', {})
        if sections.get('cover_page', False) is False:  # If cover page is included, TOC is part of it
            pass
        else:
            html += self._build_toc_html(sections)

        # Executive Summary
        if sections.get('executive_summary', True):
            html += self._build_executive_summary_html(analysis_data)

        # Data Overview
        if sections.get('data_overview', True):
            html += self._build_data_overview_html(analysis_data)

        # Data Quality
        if sections.get('data_quality', True):
            html += self._build_data_quality_html(analysis_data)

        # Summary Statistics
        if sections.get('summary_statistics', True):
            html += self._build_statistics_html(analysis_data)

        # Correlation Analysis
        if sections.get('correlation_analysis', True) and 'correlations' in analysis_data:
            html += self._build_correlation_html(analysis_data)

        # Anomalies
        if sections.get('anomalies', True) and 'anomalies' in analysis_data:
            html += self._build_anomalies_html(analysis_data)

        # Forecasts
        if sections.get('forecasts', True) and 'forecasts' in analysis_data:
            html += self._build_forecasts_html(analysis_data)

        # Conclusions
        if sections.get('conclusions', True):
            html += self._build_conclusions_html(analysis_data)

        # Footer
        html += f"""
    <div class="footer">
        <p>Generated by DataInsight Pro v1.1.0 | Enterprise Analytics Platform<br>
        Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
</body>
</html>
"""
        return html

    def _build_toc_html(self, sections: Dict) -> str:
        """Build table of contents"""
        toc_items = []
        section_names = {
            'executive_summary': 'Executive Summary',
            'data_overview': 'Data Overview',
            'data_quality': 'Data Quality Report',
            'summary_statistics': 'Summary Statistics',
            'visualizations': 'Visualizations',
            'correlation_analysis': 'Correlation Analysis',
            'anomalies': 'Anomalies Detected',
            'forecasts': 'Forecasts',
            'conclusions': 'Conclusions'
        }

        for key, name in section_names.items():
            if sections.get(key, True):
                toc_items.append(f'<li>{name}</li>')

        toc_html = f"""
<div class="section">
    <h2>📋 Table of Contents</h2>
    <ol style="font-size: 1.1rem; line-height: 2;">
        {''.join(toc_items)}
    </ol>
</div>
"""
        return toc_html

    def _build_executive_summary_html(self, data: Dict) -> str:
        """Build executive summary section"""
        quality = data.get('quality_score', {})
        overview = data.get('overview', {})

        return f"""
<div class="section">
    <h2>📊 Executive Summary</h2>
    <p>This report provides a comprehensive analysis of <strong>{data.get('file_name', 'the dataset')}</strong>.</p>

    <div class="metric-grid">
        <div class="metric-card">
            <div class="metric-label">Total Rows</div>
            <div class="metric-value">{overview.get('rows', 0):,}</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Total Columns</div>
            <div class="metric-value">{overview.get('columns', 0)}</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Data Quality</div>
            <div class="metric-value">{quality.get('overall', 0):.1f}%</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Missing Values</div>
            <div class="metric-value">{quality.get('missing_cells', 0):,}</div>
        </div>
    </div>

    <h3>Key Findings</h3>
    <ul>
        {''.join(f'<li>{finding}</li>' for finding in data.get('key_findings', ['Analysis completed successfully']))}
    </ul>
</div>
"""

    def _build_data_overview_html(self, data: Dict) -> str:
        """Build data overview section"""
        overview = data.get('overview', {})
        data_types = data.get('data_types', {})

        type_rows = ""
        for type_name, columns in data_types.items():
            if columns:
                type_rows += f"""
<tr>
    <td><strong>{type_name.title()}</strong></td>
    <td>{', '.join(columns[:10])}{' ...' if len(columns) > 10 else ''}</td>
    <td>{len(columns)}</td>
</tr>
"""

        return f"""
<div class="section">
    <h2>🔍 Data Overview</h2>
    <p><strong>File:</strong> {data.get('file_name', 'N/A')}<br>
    <strong>Size:</strong> {data.get('file_size', 'N/A')}<br>
    <strong>Format:</strong> {data.get('file_format', 'N/A')}</p>

    <h3>Data Types Summary</h3>
    <table>
        <thead>
            <tr>
                <th>Type</th>
                <th>Columns</th>
                <th>Count</th>
            </tr>
        </thead>
        <tbody>
            {type_rows}
        </tbody>
    </table>
</div>
"""

    def _build_data_quality_html(self, data: Dict) -> str:
        """Build data quality section"""
        quality = data.get('quality_score', {})

        if not quality:
            return ""

        return f"""
<div class="section">
    <h2>✅ Data Quality Report</h2>
    <p>Overall data quality score: <span class="quality-score">{quality.get('overall', 0):.1f}%</span></p>

    <div class="metric-grid">
        <div class="metric-card">
            <div class="metric-label">Completeness</div>
            <div class="metric-value">{quality.get('completeness', 0):.1f}%</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Uniqueness</div>
            <div class="metric-value">{quality.get('uniqueness', 0):.1f}%</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Consistency</div>
            <div class="metric-value">{quality.get('consistency', 0):.1f}%</div>
        </div>
        <div class="metric-card">
            <div class="metric-label">Validity</div>
            <div class="metric-value">{quality.get('validity', 0):.1f}%</div>
        </div>
    </div>

    {f'<div class="alert alert-warning"><strong>Issues Found:</strong><ul>{"".join(f"<li>{issue}</li>" for issue in quality.get("issues", []))}</ul></div>' if quality.get('issues') else '<div class="alert alert-success">No significant quality issues detected.</div>'}
</div>
"""

    def _build_statistics_html(self, data: Dict) -> str:
        """Build statistics section"""
        stats = data.get('statistics', {})

        if not stats:
            return ""

        def fmt_num(val):
            if isinstance(val, (int, float)):
                return f"{val:.2f}"
            return 'N/A'

        stats_html = ""
        for col, col_stats in list(stats.items())[:10]:  # Limit to 10 columns
            if 'mean' in col_stats:
                stats_html += f"""
<tr>
    <td><strong>{col}</strong></td>
    <td>{col_stats.get('count', 0):,}</td>
    <td>{fmt_num(col_stats.get('mean', 0))}</td>
    <td>{fmt_num(col_stats.get('std', 0))}</td>
    <td>{fmt_num(col_stats.get('min', 0))}</td>
    <td>{fmt_num(col_stats.get('max', 0))}</td>
</tr>
"""

        return f"""
<div class="section">
    <h2>📈 Summary Statistics</h2>
    <table>
        <thead>
            <tr>
                <th>Column</th>
                <th>Count</th>
                <th>Mean</th>
                <th>Std Dev</th>
                <th>Min</th>
                <th>Max</th>
            </tr>
        </thead>
        <tbody>
            {stats_html}
        </tbody>
    </table>
</div>
"""

    def _build_correlation_html(self, data: Dict) -> str:
        """Build correlation analysis section"""
        correlations = data.get('correlations', {})

        if not correlations:
            return ""

        pearson = correlations.get('pearson', {})
        strong_pairs = pearson.get('strong_pairs', [])

        def fmt_corr(val):
            if isinstance(val, (int, float)):
                return f"{val:.4f}"
            return 'N/A'

        pairs_html = ""
        for pair in strong_pairs[:10]:
            pairs_html += f"""
<tr>
    <td>{pair['var1']} ↔ {pair['var2']}</td>
    <td>{fmt_corr(pair['correlation'])}</td>
    <td>{pair['strength']}</td>
</tr>
"""

        return f"""
<div class="section">
    <h2>🔗 Correlation Analysis</h2>
    <p>Pearson, Spearman, and Kendall correlations were calculated for all numeric variables.</p>

    {f'<h3>Strong Correlations (|r| ≥ 0.7)</h3><table><thead><tr><th>Variable Pair</th><th>Correlation</th><th>Strength</th></tr></thead><tbody>{pairs_html}</tbody></table>' if strong_pairs else '<p>No strong correlations detected.</p>'}
</div>
"""

    def _build_anomalies_html(self, data: Dict) -> str:
        """Build anomalies section"""
        anomalies = data.get('anomalies', {})

        if not anomalies:
            return ""

        anomaly_html = ""
        for col, info in list(anomalies.items())[:5]:
            anomaly_html += f"""
<tr>
    <td><strong>{col}</strong></td>
    <td>{info.get('count', 0)}</td>
    <td>{info.get('percentage', 0):.2f}%</td>
</tr>
"""

        return f"""
<div class="section">
    <h2>⚠️ Anomalies Detected</h2>
    <p>Statistical methods (IQR, Z-Score, Isolation Forest) were used to identify anomalies.</p>

    <table>
        <thead>
            <tr>
                <th>Column</th>
                <th>Anomalies</th>
                <th>Percentage</th>
            </tr>
        </thead>
        <tbody>
            {anomaly_html}
        </tbody>
    </table>
</div>
"""

    def _build_forecasts_html(self, data: Dict) -> str:
        """Build forecasts section"""
        forecasts = data.get('forecasts', {})

        if not forecasts:
            return ""

        forecast_text = ""
        for key, forecast in forecasts.items():
            if 'error' not in forecast:
                forecast_text += f"""
<div>
    <h3>{key}</h3>
    <p><strong>Method:</strong> {forecast.get('method', 'N/A').title()}<br>
    <strong>Periods:</strong> {forecast.get('periods', 0)}<br>
    <strong>Next value forecast:</strong> {forecast.get('forecast_values', [0])[0]:.2f if forecast.get('forecast_values') else 'N/A'}</p>
</div>
"""

        return f"""
<div class="section">
    <h2>🔮 Forecasts</h2>
    <p>Time-series forecasting has been performed using various methods.</p>
    {forecast_text}
</div>
"""

    def _build_conclusions_html(self, data: Dict) -> str:
        """Build conclusions section"""
        conclusions = data.get('conclusions', [
            'The dataset has been thoroughly analyzed using statistical methods.',
            'Data quality is within acceptable parameters for analysis.',
            'Key patterns and trends have been identified and visualized.',
            'Recommendations based on findings are provided throughout the report.'
        ])

        return f"""
<div class="section">
    <h2>💡 Conclusions</h2>
    <ul>
        {''.join(f'<li>{c}</li>' for c in conclusions)}
    </ul>

    <div class="alert alert-info">
        <strong>Note:</strong> This report is generated automatically by DataInsight Pro.
        All findings should be validated with domain expertise before making business decisions.
    </div>
</div>
"""

    def generate_pdf_report(self, analysis_data: Dict) -> Optional[bytes]:
        """Generate PDF report using ReportLab"""
        if not REPORTLAB_AVAILABLE:
            return None

        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            # Build story
            story = []
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=self.styles['title_size'],
                textColor=colors.HexColor(self.styles['primary_color']),
                spaceAfter=20,
                alignment=TA_CENTER
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=self.styles['header_size'],
                textColor=colors.HexColor(self.styles['primary_color']),
                spaceAfter=12,
                spaceBefore=12
            )

            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=self.styles['body_size'],
                spaceAfter=8,
                alignment=TA_JUSTIFY
            )

            # Cover page
            report_title = self.customizations.get('title', 'Data Analysis Report')
            story.append(Spacer(1, 2*inch))
            story.append(Paragraph(report_title, title_style))
            story.append(Spacer(1, 0.5*inch))

            cover_info = f"""
            <b>Generated:</b> {datetime.now().strftime('%B %d, %Y')}<br/>
            <b>File:</b> {analysis_data.get('file_name', 'N/A')}<br/>
            <b>Template:</b> {self.template.title()}<br/>
            """
            story.append(Paragraph(cover_info, body_style))
            story.append(PageBreak())

            # Executive Summary
            sections = self.customizations.get('sections', {})
            if sections.get('executive_summary', True):
                story.append(Paragraph('Executive Summary', heading_style))
                quality = analysis_data.get('quality_score', {})
                overview = analysis_data.get('overview', {})

                summary_text = f"""
                This report analyzes <b>{analysis_data.get('file_name', 'the dataset')}</b>
                with {overview.get('rows', 0):,} rows and {overview.get('columns', 0)} columns.
                The overall data quality score is <b>{quality.get('overall', 0):.1f}%</b>.
                """
                story.append(Paragraph(summary_text, body_style))
                story.append(Spacer(1, 0.3*inch))

            # Data Quality
            if sections.get('data_quality', True):
                story.append(Paragraph('Data Quality Report', heading_style))
                quality = analysis_data.get('quality_score', {})

                if quality:
                    quality_data = [
                        ['Metric', 'Score'],
                        ['Completeness', f"{quality.get('completeness', 0):.1f}%"],
                        ['Uniqueness', f"{quality.get('uniqueness', 0):.1f}%"],
                        ['Consistency', f"{quality.get('consistency', 0):.1f}%"],
                        ['Validity', f"{quality.get('validity', 0):.1f}%"],
                        ['Overall', f"{quality.get('overall', 0):.1f}%"]
                    ]

                    t = Table(quality_data, colWidths=[3*inch, 2*inch])
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.styles['primary_color'])),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor(self.styles['secondary_color'])),
                        ('TEXTCOLOR', (0, -1), (-1, -1), colors.whitesmoke),
                        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.3*inch))

            # Summary Statistics
            if sections.get('summary_statistics', True):
                story.append(Paragraph('Summary Statistics', heading_style))
                stats = analysis_data.get('statistics', {})

                if stats:
                    stat_data = [['Column', 'Count', 'Mean', 'Std', 'Min', 'Max']]
                    for col, col_stats in list(stats.items())[:8]:
                        if 'mean' in col_stats:
                            stat_data.append([
                                col[:20],
                                f"{col_stats.get('count', 0):,}",
                                f"{col_stats.get('mean', 0):.2f}",
                                f"{col_stats.get('std', 0):.2f}",
                                f"{col_stats.get('min', 0):.2f}",
                                f"{col_stats.get('max', 0):.2f}"
                            ])

                    if len(stat_data) > 1:
                        t = Table(stat_data)
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.styles['primary_color'])),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
                        ]))
                        story.append(t)

            # Build PDF
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception as e:
            print(f"PDF generation error: {e}")
            return None

    def generate_docx_report(self, analysis_data: Dict) -> Optional[bytes]:
        """Generate DOCX report"""
        if not DOCX_AVAILABLE:
            return None

        try:
            doc = Document()
            report_title = self.customizations.get('title', 'Data Analysis Report')

            # Title
            title = doc.add_heading(report_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
            doc.add_paragraph(f'File: {analysis_data.get("file_name", "N/A")}')
            doc.add_paragraph(f'Template: {self.template.title()}')

            doc.add_paragraph()  # Spacer

            sections = self.customizations.get('sections', {})

            # Executive Summary
            if sections.get('executive_summary', True):
                doc.add_heading('Executive Summary', 1)
                quality = analysis_data.get('quality_score', {})
                overview = analysis_data.get('overview', {})

                doc.add_paragraph(
                    f'This report analyzes {analysis_data.get("file_name", "the dataset")} '
                    f'with {overview.get("rows", 0):,} rows and {overview.get("columns", 0)} columns. '
                    f'The overall data quality score is {quality.get("overall", 0):.1f}%.'
                )

            # Data Quality
            if sections.get('data_quality', True) and analysis_data.get('quality_score'):
                doc.add_heading('Data Quality Report', 1)
                quality = analysis_data['quality_score']

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Metric'
                hdr_cells[1].text = 'Score'

                for metric in ['completeness', 'uniqueness', 'consistency', 'validity']:
                    row = table.add_row().cells
                    row[0].text = metric.title()
                    row[1].text = f"{quality.get(metric, 0):.1f}%"

                row = table.add_row().cells
                row[0].text = 'Overall'
                row[1].text = f"{quality.get('overall', 0):.1f}%"

            # Summary Statistics
            if sections.get('summary_statistics', True) and analysis_data.get('statistics'):
                doc.add_heading('Summary Statistics', 1)
                stats = analysis_data['statistics']

                # Create table
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                hdr = table.rows[0].cells
                hdr[0].text = 'Column'
                hdr[1].text = 'Mean'
                hdr[2].text = 'Std Dev'
                hdr[3].text = 'Count'

                for col, col_stats in list(stats.items())[:10]:
                    if 'mean' in col_stats:
                        row = table.add_row().cells
                        row[0].text = col
                        row[1].text = f"{col_stats.get('mean', 0):.2f}"
                        row[2].text = f"{col_stats.get('std', 0):.2f}"
                        row[3].text = f"{col_stats.get('count', 0):,}"

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception as e:
            print(f"DOCX generation error: {e}")
            return None

    def generate_json_report(self, analysis_data: Dict) -> str:
        """Generate JSON report with all analysis results"""
        report = {
            'metadata': {
                'title': self.customizations.get('title', 'Data Analysis Report'),
                'template': self.template,
                'generated_at': datetime.now().isoformat(),
                'version': '1.1.0',
                'file_name': analysis_data.get('file_name', 'N/A')
            },
            'analysis': analysis_data
        }

        return json.dumps(report, indent=2, default=str)
